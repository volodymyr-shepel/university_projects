
import requests
import os
import sys
import matplotlib.pyplot as plt
from PIL import Image
import json
import statistics
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt,RGBColor


# link : https://www.gutenberg.org/ebooks/174

# provide the page of the book e.g: https://www.gutenberg.org/ebooks/1513
# download link for the book e.g:https://www.gutenberg.org/files/1513/1513-0.txt
def ask_for_input():
    # ask for input
    book_link = input("Provide the link to the page with desired book:")
    
    filename = input("Provide the desired filename(with file extension):")
    
    while os.path.exists(filename):
        filename = input("Provided filename already exists in the current folder.Enter another filename:")
    return book_link,filename
def download_book(): 
    
    book_link,filename = ask_for_input()
    # check if the file with provided filename already exists
    
    
    book_id = book_link.split('/')[-1] # extract book_id
    download_link = f'https://www.gutenberg.org/files/{book_id}/{book_id}-0.txt'
    response = requests.get(download_link)
    
    # check the status code
    if response.status_code == 200:
        with open(filename,"wb") as file:
            file.write(response.content)
        print("File downloaded successfully.")
        return True,filename
    else:
        print("Error")
        return False

# TASK 2
def extract_data(filename):
    number_of_lines = 56 # 56,because the header in each file consists of 56 lines
    
    
    chapter_1_content = ""
    with open(filename,encoding="utf-8") as file:
        head = [file.readline() for _ in range(number_of_lines)]
        line = file.readline()
        while(line.strip() != "CHAPTER I."):
            line = file.readline()
        
        while(line.strip() != "CHAPTER II."):
            chapter_1_content = chapter_1_content + line
            line = file.readline()
    
    # write the content of the first chapter to the file    
    with open("chapter_1.txt","w",encoding="utf-8") as f:
        f.write(chapter_1_content)
    
    title = head[10] # add as a variable(constant)
    author = head[12]
    
    return title,author

# TASK 3        
def number_of_words(filename):
    with open(filename,encoding="utf-8") as file:
        content = file.read()
        paragraphs = content.split("\n\n")     
    word_counts = []
    for paragraph in paragraphs:
        words = paragraph.split()  # Split paragraph into words
        word_count = len(words)  # Count the number of words
        word_counts.append(word_count)  # Store word count
    
    bins = range(0, max(word_counts) + 5, 5)

    plt.hist(word_counts, bins=bins, edgecolor='black')
    plt.xlabel('Word Count')
    plt.ylabel('Paragraph')
    plt.title('Number of Words in Each Paragraph (Chapter I)')
    plt.savefig('word_count_plot.png')
    
    data = {}
    data["Chapter_1"] = {}
    data["Chapter_1"]["minimal_number_of_words"] = min(word_counts)
    data["Chapter_1"]["max_number_of_words"] = max(word_counts)
    data["Chapter_1"]["mean_number_of_words"] = statistics.mean(word_counts)
    data["Chapter_1"]["number_of_paragraphs"] = len(word_counts)
    
    json_data = json.dumps(data)

    with open("stats.json", "w") as json_file:
        json_file.write(json_data)

# TASK 4

def download_image(filename,url):
    response = requests.get(url)
    if response.status_code == 200:
        
        with open(filename, "wb") as file:
            file.write(response.content)

        print("Image downloaded successfully.")
    else:
        print("Failed to download the image.")

def crop_image(filename):
    img = Image.open(filename)
    w,h = img.size
    img2 = img.crop((int(0.1 * w),int(0.1 * h),int(0.9 * w), int(0.9 * h)))
    w,h = img2.size
    new_img2 = img2.resize((int(w* 0.8),int(h * 0.8)))
    new_img2.save("resized_image.png")

def add_logo(filename1,filename2): # filename1 cropped image,filename2 logo
    base_image = Image.open(filename1)
    w,h = base_image.size
    logo = Image.open(filename2)
    logo = logo.rotate(90)
    logo = logo.resize((int(w * 0.2),int(h * 0.2)))
    
    new_image = Image.new("RGBA", base_image.size)
    new_image.paste(base_image, (0, 0))
    new_image.paste(logo, (0, 0))
    new_image.save("output_image.png")

def create_doc(author,title):
    with open("stats.json", "r") as json_file:
    # Load the JSON data
        data = json.load(json_file)
    
    doc= docx.Document()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("Title page")
    run.bold = True
    run.font.size = Pt(20)
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    run.italic = True
    run.font.size = Pt(17)
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(author)
    run.bold = True
    run.italic = True
    run.font.size = Pt(17)
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run().add_picture("output_image.png")
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = paragraph.add_run("Author of the report:Volodymyr Shepel #266617")
    run.italic = True
    run.font.size = Pt(18)
    
    doc.add_paragraph()
    doc.add_page_break()
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run("Info page")
    run.bold = True
    run.font.size = Pt(20)
    paragraph = doc.add_paragraph()
    
    paragraph.add_run().add_picture("word_count_plot.png")
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for key, value in data["Chapter_1"].items():
        paragraph = doc.add_paragraph()
        paragraph.style = doc.styles['Normal']
        
        run = paragraph.add_run(key + ": " + str(value))
        
        run.bold = True
        run.italic = True
        run.font.size = Pt(15)
    doc.save("report.docx")
def main():
    folder_name = "materials"
    os.makedirs(folder_name, exist_ok=True)
    os.chdir(folder_name)
    is_successfully_downloaded,filename = download_book()
    if not is_successfully_downloaded:
        sys.exit("Failed to download the file.Try again later")
    title,author = extract_data("book.txt")
    number_of_words("chapter_1.txt")
    
    
    
    # download image
    download_image("base_image.png","https://cdn.dribbble.com/users/3217740/screenshots/16419660/media/75d20092eb684495f81ef52bab9a2237.png?compress=1&resize=400x300")
    crop_image("base_image.png") # will create resized_image.png
    
    # download logo
    download_image("logo.png","https://static.wixstatic.com/media/617b77_2643708195a0447381ab576c951a5113~mv2.jpg/v1/fill/w_1288,h_936,al_c,q_85/617b77_2643708195a0447381ab576c951a5113~mv2.jpg")
    add_logo("resized_image.png","logo.png") # will create output_image.png
    
    create_doc(author=author,title=title)
    
    # print("Done")
if __name__ == "__main__":
    main()